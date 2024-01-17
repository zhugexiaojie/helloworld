import xlrd
from docx import Document
from docx.shared import Cm, Pt
from PyQt5.QtWidgets import QApplication, QFileDialog, QWidget, QPushButton, QVBoxLayout, QMessageBox
from PyQt5.QtGui import QIcon

def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.element.clear()

class Window(QWidget):
    def __init__(self):
        super().__init__()

        self.title = "Word生成程序"
        self.width = 400
        self.height = 140
        self.initUI()

    def initUI(self):
        self.setWindowTitle(self.title)
        self.resize(self.width, self.height)
        self.move((QApplication.desktop().screenGeometry().width() - self.width) // 2,
                  (QApplication.desktop().screenGeometry().height() - self.height) // 2)

        layout = QVBoxLayout()

        self.label = QPushButton("选择Excel文件", self)
        layout.addWidget(self.label)

        self.button = QPushButton('选择Excel文件', self)
        self.button.clicked.connect(self.generate)
        self.button.setIcon(QIcon('excel_icon.png'))
        self.button.setStyleSheet("QPushButton {"
                                   "background-color: #4CAF50;"
                                   "border: none;"
                                   "color: white;"
                                   "padding: 8px 12px;"
                                   "text-align: center;"
                                   "text-decoration: none;"
                                   "display: inline-block;"
                                   "font-size: 9px;"
                                 "}")
        layout.addWidget(self.button)

        self.setLayout(layout)
        
    def generate(self):
        excel_path, _ = QFileDialog.getOpenFileName(self, "选择Excel文件", "", "Excel文件(*.xls *.xlsx)")  
        if excel_path:
            if excel_path.endswith('.xls'):
                wb = xlrd.open_workbook(excel_path)
                sheet = wb.sheet_by_index(0)  # 假设只有一个工作表，你可以根据需要修改索引
                data = [sheet.cell_value(row, 1) for row in range(1, sheet.nrows) if sheet.cell_value(row, 1)]

                doc = Document() 
                total_cells = 10 * 6
                total_pages = (len(data) + total_cells - 1) // total_cells

                for page in range(total_pages):
                    table = doc.add_table(rows=10, cols=6)
                    table.autofit = False
                    for row in table.rows:
                        row.height = Cm(2.76)
                    remove_table_borders(table)
                    for r in range(10):
                        for c in range(6):
                            index = page * total_cells + r * 6 + c
                            if index < len(data):
                                cell = table.cell(r, c)
                                cell.text = '重做\n%s' % data[index]
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = '宋体'
                                        run.font.size = Pt(9)
                                        paragraph.alignment = 1
                                        cell.vertical_alignment = 1
                                        run.font.size = Pt(9)

                save_path, _ = QFileDialog.getSaveFileName(self, "保存Word文件", "", "Word文件 (*.docx)")
                if save_path:
                    doc.save(save_path)
                    QMessageBox.information(self, "完成", "Word文件已保存至：%s" % save_path, QMessageBox.Ok)
        
app = QApplication([])
window = Window()
window.show()
app.exec()
